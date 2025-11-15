#!/usr/bin/env python3
"""
Flask OCR Web Application
Giao diện web cho nhận dạng văn bản tiếng Việt
"""
from flask import Flask, render_template, request, jsonify, send_file, flash, redirect, url_for
import cv2
import numpy as np
from PIL import Image
import json
import os
import tempfile
from datetime import datetime
from pathlib import Path
import time
import io
import base64
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import zipfile

# Sử dụng model_loader và utils đã cải thiện
from model_loader import ModelLoader
from utils import preprocess_image, decode_predictions

# Sử dụng OCR không cần train (EasyOCR/PaddleOCR)
try:
    from ocr_without_training import OCRWithoutTraining
    OCR_WITHOUT_TRAINING_AVAILABLE = True
except ImportError:
    OCR_WITHOUT_TRAINING_AVAILABLE = False
    print("⚠️  OCRWithoutTraining không khả dụng. Chỉ sử dụng model đã train.")

# Sử dụng PaddleOCR cho văn bản dài
try:
    from paddleocr_long_text import PaddleOCRLongText
    PADDLEOCR_LONG_TEXT_AVAILABLE = True
except ImportError:
    PADDLEOCR_LONG_TEXT_AVAILABLE = False
    print("⚠️  PaddleOCRLongText không khả dụng.")

# Sử dụng PaddleOCR cho chữ viết tay
try:
    from paddleocr_handwritten import PaddleOCRHandwritten
    PADDLEOCR_HANDWRITTEN_AVAILABLE = True
except ImportError:
    PADDLEOCR_HANDWRITTEN_AVAILABLE = False
    print("⚠️  PaddleOCRHandwritten không khả dụng.")

# Sử dụng Combined OCR (CRNN + PaddleOCR)
try:
    from ocr_combined_crnn_paddle import CombinedOCR
    COMBINED_OCR_AVAILABLE = True
except ImportError:
    COMBINED_OCR_AVAILABLE = False
    print("⚠️  CombinedOCR không khả dụng.")

app = Flask(__name__)
app.secret_key = 'your-secret-key-here'  # Thay đổi key này trong production

# Add static file serving for sample images
@app.route('/data/<path:filename>')
def serve_data(filename):
    """Serve data files including sample images"""
    return send_file(f'data/{filename}')

# Global variables for model
model_loader = None
ocr_engine = None
paddleocr_long = None  # PaddleOCR cho văn bản dài
paddleocr_handwritten = None  # PaddleOCR cho chữ viết tay
combined_ocr = None  # Combined OCR (CRNN + PaddleOCR)
use_trained_model = True  # True: dùng model đã train, False: dùng OCR không cần train
ocr_mode = 'auto'  # 'auto', 'crnn', 'paddleocr', 'paddleocr_handwritten', 'easyocr', 'combined'

def load_ocr_model():
    """Load OCR model using improved ModelLoader"""
    global model_loader
    try:
        print("Đang tải model...")
        model_loader = ModelLoader()
        print("✅ OCR model loaded successfully!")
        return True
    except Exception as e:
        print(f"❌ Error loading OCR model: {e}")
        import traceback
        traceback.print_exc()
        return False

def init_ocr_without_training(method='easyocr'):
    """Initialize OCR engine không cần train"""
    global ocr_engine
    if not OCR_WITHOUT_TRAINING_AVAILABLE:
        return False
    try:
        print(f"Đang khởi tạo OCR engine: {method}...")
        ocr_engine = OCRWithoutTraining(method=method)
        print("✅ OCR engine đã sẵn sàng")
        return True
    except Exception as e:
        print(f"❌ Lỗi khởi tạo OCR engine: {e}")
        return False

def init_paddleocr_long_text():
    """Initialize PaddleOCR cho văn bản dài"""
    global paddleocr_long
    if not PADDLEOCR_LONG_TEXT_AVAILABLE:
        return False
    try:
        print("Đang khởi tạo PaddleOCR cho văn bản dài...")
        paddleocr_long = PaddleOCRLongText(lang='vi')
        print("✅ PaddleOCR cho văn bản dài đã sẵn sàng")
        return True
    except Exception as e:
        print(f"❌ Lỗi khởi tạo PaddleOCR: {e}")
        return False

def init_paddleocr_handwritten():
    """Initialize PaddleOCR cho chữ viết tay"""
    global paddleocr_handwritten
    if not PADDLEOCR_HANDWRITTEN_AVAILABLE:
        return False
    try:
        print("Đang khởi tạo PaddleOCR cho chữ viết tay...")
        paddleocr_handwritten = PaddleOCRHandwritten(lang='vi')
        print("✅ PaddleOCR cho chữ viết tay đã sẵn sàng")
        return True
    except Exception as e:
        print(f"❌ Lỗi khởi tạo PaddleOCR Handwritten: {e}")
        return False

def init_combined_ocr():
    """Initialize Combined OCR (CRNN + PaddleOCR)"""
    global combined_ocr
    if not COMBINED_OCR_AVAILABLE:
        return False
    try:
        print("Đang khởi tạo Combined OCR (CRNN + PaddleOCR)...")
        combined_ocr = CombinedOCR(use_paddle=True)
        print("✅ Combined OCR đã sẵn sàng")
        return True
    except Exception as e:
        print(f"❌ Lỗi khởi tạo Combined OCR: {e}")
        import traceback
        traceback.print_exc()
        return False

def predict_image(image, mode=None):
    """
    Predict text from image
    
    Args:
        image: PIL Image
        mode: 'auto', 'crnn', 'paddleocr', 'paddleocr_handwritten', 'easyocr', 'combined'
    """
    global use_trained_model, ocr_engine, paddleocr_long, paddleocr_handwritten, combined_ocr, ocr_mode
    
    # Sử dụng mode được chỉ định hoặc mode mặc định
    current_mode = mode if mode else ocr_mode
    
    # Mode: combined - Sử dụng Combined OCR (CRNN + PaddleOCR)
    if current_mode == 'combined':
        if combined_ocr is not None and hasattr(combined_ocr, 'recognize'):
            try:
                result = combined_ocr.recognize(image, method='combined')
                if result and result.get('text', '').strip():
                    return result['text']
            except Exception as e:
                print(f"Error with Combined OCR: {e}")
                import traceback
                traceback.print_exc()
        else:
            print("⚠️  Combined OCR không khả dụng")
    
    # Mode: paddleocr_handwritten - Sử dụng PaddleOCR cho chữ viết tay
    if current_mode == 'paddleocr_handwritten':
        if paddleocr_handwritten is not None and hasattr(paddleocr_handwritten, 'recognize_with_boxes'):
            try:
                result = paddleocr_handwritten.recognize_with_boxes(image, return_image=False)
                if result and result.get('text', '').strip():
                    return result['text']
            except Exception as e:
                print(f"Error with PaddleOCR Handwritten: {e}")
                import traceback
                traceback.print_exc()
        else:
            print("⚠️  PaddleOCR Handwritten không khả dụng")
    
    # Mode: paddleocr - Sử dụng PaddleOCR cho văn bản dài
    if current_mode == 'paddleocr':
        if paddleocr_long is not None and hasattr(paddleocr_long, 'recognize_with_layout'):
            try:
                result = paddleocr_long.recognize_with_layout(image)
                if result and result.get('text', '').strip():
                    return result['text']
            except Exception as e:
                print(f"Error with PaddleOCR: {e}")
                import traceback
                traceback.print_exc()
        else:
            print("⚠️  PaddleOCR Long không khả dụng")
    
    # Mode: easyocr hoặc auto - Sử dụng EasyOCR
    if (current_mode == 'easyocr' or current_mode == 'auto'):
        # Kiểm tra ocr_engine có tồn tại và có method recognize không
        if ocr_engine is not None and hasattr(ocr_engine, 'recognize'):
            try:
                predicted_text = ocr_engine.recognize(image)
                if predicted_text and predicted_text.strip():
                    return predicted_text
                else:
                    # Nếu EasyOCR không trả về kết quả, thử fallback
                    print("⚠️  EasyOCR không trả về kết quả, thử engine khác...")
            except Exception as e:
                print(f"Error with OCR engine: {e}")
                import traceback
                traceback.print_exc()
                # Fallback sang engine khác nếu có
                if paddleocr_handwritten is not None:
                    try:
                        print("🔄 Thử PaddleOCR Handwritten...")
                        result = paddleocr_handwritten.recognize_with_boxes(image, return_image=False)
                        if result and result.get('text', '').strip():
                            return result['text']
                    except:
                        pass
        else:
            print("⚠️  OCR engine không khả dụng, thử engine khác...")
    
    # Mode: crnn hoặc fallback - Sử dụng model đã train
    if (current_mode == 'crnn' or current_mode == 'auto') and use_trained_model:
        # Kiểm tra model_loader có tồn tại và có method predict không
        if model_loader is not None and hasattr(model_loader, 'predict'):
            try:
                # Convert PIL image to bytes for preprocessing
                img_bytes = io.BytesIO()
                if image.mode != 'RGB':
                    image = image.convert('RGB')
                image.save(img_bytes, format='PNG')
                img_bytes.seek(0)
                img_bytes = img_bytes.read()
                
                # Preprocess image using improved preprocessing (khớp với training)
                processed_image = preprocess_image(img_bytes)
                
                # Verify processed image
                if processed_image is None or processed_image.size == 0:
                    raise ValueError("Processed image is empty")
                
                # Verify shape matches model input
                expected_shape = (1, 118, 2167, 1)
                if processed_image.shape != expected_shape:
                    raise ValueError(f"Image shape {processed_image.shape} doesn't match expected {expected_shape}")
                
                # Predict using model_loader
                predictions = model_loader.predict(processed_image)
                
                # Decode predictions
                decoded_text = decode_predictions(predictions, model_loader.char_list, greedy=True)
                
                return decoded_text
            except Exception as e:
                import traceback
                error_msg = f"Error during prediction with trained model: {e}"
                print(error_msg)
                print(traceback.format_exc())
        else:
            print("⚠️  Model loader không khả dụng, bỏ qua CRNN mode")
    
    # Nếu cả 2 đều fail - thử tất cả các engine còn lại
    print("⚠️  Tất cả engine chính đã fail, thử các engine backup...")
    
    # Thử PaddleOCR Handwritten nếu chưa thử
    if paddleocr_handwritten is not None and current_mode != 'paddleocr_handwritten':
        try:
            result = paddleocr_handwritten.recognize_with_boxes(image, return_image=False)
            if result and result.get('text', '').strip():
                return result['text']
        except Exception as e:
            print(f"PaddleOCR Handwritten failed: {e}")
    
    # Thử PaddleOCR Long nếu chưa thử
    if paddleocr_long is not None and current_mode != 'paddleocr':
        try:
            result = paddleocr_long.recognize_with_layout(image)
            if result and result.get('text', '').strip():
                return result['text']
        except Exception as e:
            print(f"PaddleOCR Long failed: {e}")
    
    # Nếu vẫn fail
    raise ValueError("Không thể nhận dạng ảnh. Vui lòng kiểm tra:\n- Ảnh có chất lượng tốt không\n- Ảnh có chứa chữ viết tay không\n- Thử ảnh khác hoặc chọn mode khác")


def create_txt_file(text, filename_prefix="ocr_result"):
    """Create TXT file"""
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    temp_dir = tempfile.mkdtemp()
    txt_path = os.path.join(temp_dir, f"{filename_prefix}_{timestamp}.txt")
    
    with open(txt_path, 'w', encoding='utf-8') as f:
        f.write(text)
    
    return txt_path

def create_word_file(text, filename_prefix="ocr_result"):
    """Create Word document"""
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    temp_dir = tempfile.mkdtemp()
    docx_path = os.path.join(temp_dir, f"{filename_prefix}_{timestamp}.docx")
    
    # Create new document
    doc = Document()
    
    # Add title
    title = doc.add_heading('Kết quả nhận dạng OCR', 0)
    
    # Add timestamp
    doc.add_paragraph(f'Thời gian tạo: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    doc.add_paragraph('')
    
    # Add content
    doc.add_paragraph(text)
    
    # Save document
    doc.save(docx_path)
    
    return docx_path

def create_pdf_file(text, filename_prefix="ocr_result"):
    """Create PDF file"""
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    temp_dir = tempfile.mkdtemp()
    pdf_path = os.path.join(temp_dir, f"{filename_prefix}_{timestamp}.pdf")
    
    # Create PDF document
    doc = SimpleDocTemplate(pdf_path, pagesize=letter)
    styles = getSampleStyleSheet()
    story = []
    
    # Title
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=16,
        spaceAfter=30,
        alignment=1  # Center alignment
    )
    story.append(Paragraph("Kết quả nhận dạng OCR", title_style))
    story.append(Spacer(1, 12))
    
    # Timestamp
    story.append(Paragraph(f"Thời gian tạo: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", styles['Normal']))
    story.append(Spacer(1, 12))
    
    # Content
    # Split text into paragraphs for better formatting
    paragraphs = text.split('\n')
    for para in paragraphs:
        if para.strip():
            story.append(Paragraph(para, styles['Normal']))
            story.append(Spacer(1, 6))
    
    # Build PDF
    doc.build(story)
    
    return pdf_path

@app.route('/')
def index():
    """Home"""
    return render_template('index.html')



@app.route('/upload', methods=['POST'])
def upload_file():
    """Xử lý upload file"""
    global ocr_mode
    try:
        # Lấy mode từ request (nếu có)
        mode = request.form.get('ocr_mode', ocr_mode)
        # Check if file exists in request
        if 'file' not in request.files:
            return jsonify({'error': 'Không có file được chọn'}), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': 'Không có file được chọn'}), 400
        
        # Check file extension
        if not allowed_file(file.filename):
            return jsonify({'error': f'Định dạng file không được hỗ trợ. Hỗ trợ: PNG, JPG, JPEG, GIF, BMP'}), 400
        
        # Check file size (max 10MB)
        file.seek(0, 2)  # Seek to end
        file_size = file.tell()
        file.seek(0)  # Reset to beginning
        
        if file_size > 10 * 1024 * 1024:  # 10MB
            return jsonify({'error': 'File quá lớn. Kích thước tối đa: 10MB'}), 400
        
        # Read and validate image
        try:
            file.stream.seek(0)  # Reset stream position
            image = Image.open(file.stream)
            # Verify image
            image.verify()
            # Reopen image after verify (verify closes the image)
            file.stream.seek(0)
            image = Image.open(file.stream)
            
            # Convert to RGB if needed
            if image.mode != 'RGB':
                image = image.convert('RGB')
            
            # Verify image dimensions
            if image.size[0] == 0 or image.size[1] == 0:
                return jsonify({'error': 'Ảnh không hợp lệ: kích thước bằng 0'}), 400
                
        except Exception as e:
            import traceback
            print(f"Error reading image: {e}")
            print(traceback.format_exc())
            return jsonify({'error': f'Không thể đọc file ảnh: {str(e)}'}), 400
        
        # Process with OCR
        start_time = time.time()
        try:
            predicted_text = predict_image(image, mode=mode)
            processing_time = time.time() - start_time
            
            # Check if prediction failed
            if not predicted_text or predicted_text.strip() == '':
                return jsonify({
                    'error': 'Không thể nhận dạng văn bản từ ảnh. Vui lòng thử với ảnh khác hoặc kiểm tra chất lượng ảnh.',
                    'processing_time': round(processing_time, 2)
                }), 400
            
            # Calculate confidence - Ưu tiên EasyOCR
            if ocr_engine is not None and hasattr(ocr_engine, 'get_confidence'):
                # Ưu tiên EasyOCR (đã khởi tạo thành công)
                confidence = ocr_engine.get_confidence()
                if confidence == 0:
                    confidence = 90.0  # Default
            elif mode == 'paddleocr_handwritten' and paddleocr_handwritten is not None and hasattr(paddleocr_handwritten, 'get_confidence'):
                # Lấy confidence từ PaddleOCR Handwritten
                confidence = paddleocr_handwritten.get_confidence()
                if confidence == 0:
                    confidence = 90.0  # Default
            elif mode == 'paddleocr' and paddleocr_long is not None and hasattr(paddleocr_long, 'get_confidence'):
                # Lấy confidence từ PaddleOCR
                confidence = paddleocr_long.get_confidence()
                if confidence == 0:
                    confidence = 90.0  # Default
            elif use_trained_model and model_loader is not None and hasattr(model_loader, 'predict'):
                # Calculate confidence from prediction (model đã train) - chỉ khi có model
                try:
                    from utils import calculate_confidence
                    # Get prediction for confidence calculation
                    img_bytes = io.BytesIO()
                    if image.mode != 'RGB':
                        image = image.convert('RGB')
                    image.save(img_bytes, format='PNG')
                    img_bytes.seek(0)
                    processed_img = preprocess_image(img_bytes.read())
                    prediction = model_loader.predict(processed_img)
                    confidence = calculate_confidence(prediction)
                except Exception as e:
                    print(f"⚠️  Không thể tính confidence từ model: {e}")
                    confidence = 85.0  # Default fallback
            else:
                # Default confidence nếu không có engine nào
                confidence = 85.0
            
            # Convert numpy types to Python native types for JSON serialization
            import numpy as np
            if isinstance(confidence, (np.floating, np.integer)):
                confidence = float(confidence)
            if isinstance(processing_time, (np.floating, np.integer)):
                processing_time = float(processing_time)
            
            result = {
                'text': predicted_text,
                'processing_time': round(float(processing_time), 2),
                'confidence': round(float(confidence), 1),
                'timestamp': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
                'char_count': len(predicted_text),
                'ocr_mode': mode
            }
            
            return jsonify(result)
        except Exception as e:
            import traceback
            processing_time = time.time() - start_time
            error_msg = f'Lỗi xử lý OCR: {str(e)}'
            print(error_msg)
            print(traceback.format_exc())
            return jsonify({
                'error': error_msg,
                'processing_time': round(processing_time, 2)
            }), 500
        
    except Exception as e:
        print(f"Upload error: {e}")
        return jsonify({'error': f'Lỗi xử lý: {str(e)}'}), 500

@app.route('/download/<file_type>')
def download_file(file_type):
    """Tải file về"""
    text = request.args.get('text', '')
    if not text:
        return jsonify({'error': 'Không có nội dung để tải'}), 400
    
    try:
        if file_type == 'txt':
            file_path = create_txt_file(text)
            return send_file(file_path, as_attachment=True, 
                           download_name=f"ocr_result_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt")
        
        elif file_type == 'docx':
            file_path = create_word_file(text)
            return send_file(file_path, as_attachment=True,
                           download_name=f"ocr_result_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx")
        
        elif file_type == 'pdf':
            file_path = create_pdf_file(text)
            return send_file(file_path, as_attachment=True,
                           download_name=f"ocr_result_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf")
        
        else:
            return jsonify({'error': 'Định dạng file không được hỗ trợ'}), 400
            
    except Exception as e:
        return jsonify({'error': f'Lỗi tạo file: {str(e)}'}), 500

@app.route('/download_all')
def download_all():
    """Tải tất cả định dạng file"""
    text = request.args.get('text', '')
    if not text:
        return jsonify({'error': 'Không có nội dung để tải'}), 400
    
    try:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        temp_dir = tempfile.mkdtemp()
        zip_path = os.path.join(temp_dir, f"ocr_result_{timestamp}.zip")
        
        # Create all file types
        txt_path = create_txt_file(text, "ocr_result")
        docx_path = create_word_file(text, "ocr_result")
        pdf_path = create_pdf_file(text, "ocr_result")
        
        # Create ZIP file
        with zipfile.ZipFile(zip_path, 'w') as zipf:
            zipf.write(txt_path, "ocr_result.txt")
            zipf.write(docx_path, "ocr_result.docx")
            zipf.write(pdf_path, "ocr_result.pdf")
        
        return send_file(zip_path, as_attachment=True,
                        download_name=f"ocr_result_{timestamp}.zip")
        
    except Exception as e:
        return jsonify({'error': f'Lỗi tạo file ZIP: {str(e)}'}), 500

def allowed_file(filename):
    """Kiểm tra định dạng file được phép"""
    ALLOWED_EXTENSIONS = {'png', 'jpg', 'jpeg', 'gif', 'bmp', 'PNG', 'JPG', 'JPEG', 'GIF', 'BMP'}
    return '.' in filename and filename.rsplit('.', 1)[1] in ALLOWED_EXTENSIONS

@app.route('/samples')
def samples():
    """Trang ảnh mẫu"""
    return render_template('samples.html')

@app.route('/info')
def info():
    """Trang thông tin hệ thống"""
    return render_template('info.html')


if __name__ == '__main__':
    # Khởi tạo các OCR engines
    ocr_engine_loaded = False
    paddleocr_loaded = False
    model_loaded = False
    combined_ocr_loaded = False
    paddleocr_handwritten_loaded = False
    
    # 0. Khởi tạo EasyOCR - ưu tiên cao nhất
    if OCR_WITHOUT_TRAINING_AVAILABLE:
        print("🔄 Đang khởi tạo EasyOCR...")
        ocr_engine_loaded = init_ocr_without_training('easyocr')
        if ocr_engine_loaded:
            ocr_mode = 'easyocr'  # Mặc định dùng EasyOCR
            print("✅ Sử dụng EasyOCR")
    
    # 1. Khởi tạo Combined OCR (CRNN + PaddleOCR) - backup
    if COMBINED_OCR_AVAILABLE and not ocr_engine_loaded:
        print("🔄 Đang khởi tạo Combined OCR (CRNN + PaddleOCR)...")
        combined_ocr_loaded = init_combined_ocr()
        if combined_ocr_loaded:
            ocr_mode = 'combined'  # Dùng Combined OCR nếu EasyOCR không có
            print("✅ Sử dụng Combined OCR (CRNN + PaddleOCR)")
    
    # 2. Khởi tạo PaddleOCR cho chữ viết tay (backup)
    if PADDLEOCR_HANDWRITTEN_AVAILABLE and not ocr_engine_loaded and not combined_ocr_loaded:
        print("🔄 Đang khởi tạo PaddleOCR cho chữ viết tay...")
        paddleocr_handwritten_loaded = init_paddleocr_handwritten()
        if paddleocr_handwritten_loaded:
            ocr_mode = 'paddleocr_handwritten'  # Mặc định dùng PaddleOCR Handwritten
            print("✅ Sử dụng PaddleOCR cho chữ viết tay")
    
    # 3. Khởi tạo PaddleOCR cho văn bản dài
    if PADDLEOCR_LONG_TEXT_AVAILABLE:
        print("🔄 Đang khởi tạo PaddleOCR cho văn bản dài...")
        paddleocr_loaded = init_paddleocr_long_text()
        if paddleocr_loaded and not paddleocr_handwritten_loaded:
            ocr_mode = 'paddleocr'  # Dùng PaddleOCR nếu không có Handwritten
            print("✅ Sử dụng PaddleOCR cho văn bản dài")
    
    # 4. Khởi tạo EasyOCR/PaddleOCR thông thường (nếu chưa có)
    if OCR_WITHOUT_TRAINING_AVAILABLE and not ocr_engine_loaded:
        print("🔄 Đang khởi tạo OCR engine (EasyOCR/PaddleOCR)...")
        methods = ['easyocr', 'paddleocr', 'tesseract']
        for method in methods:
            if init_ocr_without_training(method):
                ocr_engine_loaded = True
                print(f"✅ OCR engine ({method}) đã sẵn sàng")
                break
    
    # 5. Load model CRNN đã train (cho chữ viết tay - backup)
    if not ocr_engine_loaded and not combined_ocr_loaded:
        print("🔄 Đang load model CRNN cho chữ viết tay...")
        model_loaded = load_ocr_model()
        if model_loaded:
            print("✅ Model CRNN đã sẵn sàng")
    
    # Khởi động Flask app
    if combined_ocr_loaded or paddleocr_handwritten_loaded or paddleocr_loaded or ocr_engine_loaded or model_loaded:
        print("\n🚀 Starting Flask OCR Application...")
        print("=" * 70)
        print("CÁC ENGINE ĐÃ SẴN SÀNG:")
        if ocr_engine_loaded:
            print("   ✅ EasyOCR - Mode: easyocr (Mặc định)")
        if combined_ocr_loaded:
            print("   ✅ Combined OCR (CRNN + PaddleOCR) - Mode: combined")
        if paddleocr_handwritten_loaded:
            print("   ✅ PaddleOCR Handwritten (chữ viết tay) - Mode: paddleocr_handwritten")
        if paddleocr_loaded:
            print("   ✅ PaddleOCR (văn bản dài) - Mode: paddleocr")
        if ocr_engine_loaded:
            print("   ✅ OCR Engine (EasyOCR/PaddleOCR) - Mode: easyocr")
        if model_loaded:
            print("   ✅ CRNN Model (chữ viết tay - backup) - Mode: crnn")
        print("=" * 70)
        print(f"   Mode mặc định: {ocr_mode}")
        print("   Có thể thay đổi mode trong request (ocr_mode parameter)")
        print("=" * 70)
        app.run(debug=True, host='0.0.0.0', port=5000)
    else:
        print("❌ Không thể khởi động hệ thống OCR.")
        print("   Vui lòng cài đặt:")
        print("   1. PaddleOCR: pip install paddlepaddle paddleocr")
        print("   2. EasyOCR: pip install easyocr")
        print("   3. Hoặc kiểm tra model files (model_checkpoint_weights.hdf5)")
